#!/usr/bin/env python3
"""
PDF Conversion Service
Handles PDF to image, text, and DOCX conversions
"""

import sys
import os
import io
from pathlib import Path

def pdf_to_image(input_path: str, output_path: str, output_format: str, page_num: int = 0):
    """Convert PDF to image using PyMuPDF (fitz)"""
    try:
        import fitz  # PyMuPDF
        from PIL import Image
        
        # Open PDF
        pdf_doc = fitz.open(input_path)
        
        # Convert first page (or specified page)
        page = pdf_doc[page_num] if page_num < len(pdf_doc) else pdf_doc[0]
        
        # Render page to image with high DPI
        mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
        pix = page.get_pixmap(matrix=mat)
        
        # Convert to PIL Image
        img_data = pix.tobytes("ppm")
        img = Image.open(io.BytesIO(img_data))
        
        # Save with optimization
        if output_format.upper() in ['JPG', 'JPEG']:
            # Convert RGBA to RGB for JPEG
            if img.mode == 'RGBA':
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1] if len(img.split()) == 4 else None)
                img = background
            img.save(output_path, 'JPEG', quality=85, optimize=True, progressive=True)
        elif output_format.upper() == 'PNG':
            img.save(output_path, 'PNG', optimize=True, compress_level=9)
        else:
            img.save(output_path, output_format.upper(), optimize=True)
        
        pdf_doc.close()
        return True
        
    except ImportError:
        # Fallback to pdf2image
        try:
            from pdf2image import convert_from_path
            from PIL import Image
            
            # Convert PDF to images
            images = convert_from_path(input_path, first_page=page_num+1, last_page=page_num+1, dpi=200)
            
            if images:
                img = images[0]
                
                # Save with optimization
                if output_format.upper() in ['JPG', 'JPEG']:
                    if img.mode == 'RGBA':
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[-1] if len(img.split()) == 4 else None)
                        img = background
                    img.save(output_path, 'JPEG', quality=85, optimize=True, progressive=True)
                elif output_format.upper() == 'PNG':
                    img.save(output_path, 'PNG', optimize=True, compress_level=9)
                else:
                    img.save(output_path, output_format.upper(), optimize=True)
                return True
            return False
            
        except ImportError:
            raise Exception("PDF to image conversion requires PyMuPDF or pdf2image library")
    
    except Exception as e:
        raise Exception(f"PDF to image conversion failed: {e}")

def pdf_to_text(input_path: str, output_path: str):
    """Extract text from PDF"""
    try:
        import fitz  # PyMuPDF
        
        # Open PDF
        pdf_doc = fitz.open(input_path)
        
        # Extract text from all pages
        text_content = []
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            text = page.get_text()
            text_content.append(f"Page {page_num + 1}:\n{text}\n\n")
        
        # Write to text file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))
        
        pdf_doc.close()
        return True
        
    except ImportError:
        # Fallback to pdfplumber
        try:
            import pdfplumber
            
            text_content = []
            with pdfplumber.open(input_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"Page {i + 1}:\n{text}\n\n")
            
            # Write to text file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            return True
            
        except ImportError:
            raise Exception("PDF to text conversion requires PyMuPDF or pdfplumber library")
    
    except Exception as e:
        raise Exception(f"PDF to text conversion failed: {e}")

def pdf_to_docx(input_path: str, output_path: str):
    """Convert PDF to DOCX"""
    try:
        from pdf2docx import Converter
        
        # Convert PDF to DOCX
        cv = Converter(input_path)
        cv.convert(output_path, start=0)
        cv.close()
        return True
        
    except ImportError:
        # Fallback: Extract text and create simple DOCX
        try:
            import fitz  # PyMuPDF
            from docx import Document
            
            # Extract text using PyMuPDF
            pdf_doc = fitz.open(input_path)
            doc = Document()
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                text = page.get_text()
                
                # Add page heading
                doc.add_heading(f'Page {page_num + 1}', level=2)
                
                # Add text content
                if text.strip():
                    doc.add_paragraph(text)
                else:
                    doc.add_paragraph('[No text content found on this page]')
                
                # Add page break except for last page
                if page_num < len(pdf_doc) - 1:
                    doc.add_page_break()
            
            doc.save(output_path)
            pdf_doc.close()
            return True
            
        except ImportError:
            raise Exception("PDF to DOCX conversion requires pdf2docx or python-docx with PyMuPDF")
    
    except Exception as e:
        raise Exception(f"PDF to DOCX conversion failed: {e}")

def main():
    if len(sys.argv) < 4:
        print("ERROR: Insufficient arguments")
        print("Usage: python pdf_converter.py <input_path> <output_path> <output_format>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    output_format = sys.argv[3].lower()
    
    try:
        # Import here to avoid import errors in main conversion flow
        import io
        
        if output_format in ['jpg', 'jpeg', 'png']:
            success = pdf_to_image(input_path, output_path, output_format)
        elif output_format == 'txt':
            success = pdf_to_text(input_path, output_path)
        elif output_format == 'docx':
            success = pdf_to_docx(input_path, output_path)
        else:
            raise Exception(f"Unsupported output format: {output_format}")
        
        if success:
            print("SUCCESS")
        else:
            print("ERROR: Conversion failed")
            
    except Exception as e:
        print(f"ERROR: {e}")

if __name__ == "__main__":
    main()